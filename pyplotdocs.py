import matplotlib.pyplot as plt
import networkx as nx
from docx import Document
import numpy as np
from docx.shared import Inches

# Create your directed graph
direction = nx.DiGraph()

# Nodes for NIST construct
nist_nodes = ['Govern', 'Identity', 'Protect', 'Detect', 'Respond', 'Recover']

# Nodes for MediaPro contracts
mediapro_nodes = ['Analyze', 'Plan', 'Train', 'Reinforce']

# Adding the nodes to the graph
direction.add_nodes_from(nist_nodes)
direction.add_nodes_from(mediapro_nodes)

# Edges dependencies
nist_edges = [
    ('Govern', 'Identity'), ('Govern', 'Protect'), ('Govern', 'Detect'),
    ('Govern', 'Respond'), ('Govern', 'Recover'), ('Identity', 'Protect'),
    ('Protect', 'Detect'), ('Detect', 'Respond'), ('Respond', 'Recover'),
    ('Recover', 'Identity')
]

# MediaPro edges
mediaPro_edges = [
    ('Analyze', 'Plan'), ('Plan', 'Train'), ('Train', 'Reinforce'),
    ('Reinforce', 'Analyze')
]

# Integration edges
integrate_edges = [
    ('Analyze', 'Identity'), ('Plan', 'Protect'), ('Train', 'Detect'),
    ('Reinforce', 'Respond'), ('Analyze', 'Recover')
]

# Adding edges to the graph
direction.add_edges_from(nist_edges + mediaPro_edges + integrate_edges)

# Define positions
positions = {}

# Position for the center node
positions['Govern'] = (0, 0)

# Number of nodes to place in the circle
num_nodes = len(nist_nodes) + len(mediapro_nodes) - 1  # excluding the center node

# Radius of the circle
radius = 3

# Angle step
angle_step = 2 * np.pi / num_nodes

# Place the nodes in a circle
for i, node in enumerate(nist_nodes + mediapro_nodes):
    if node == 'Govern':
        continue
    angle = i * angle_step
    x = radius * np.cos(angle)
    y = radius * np.sin(angle)
    positions[node] = (x, y)

# Create a new Word document
doc = Document()

# Add a title to the document
doc.add_heading('Graph Visualization', level=1)

# Create a figure and draw your graph
plt.figure(figsize=(6, 4))
nx.draw(direction, pos=positions, with_labels=True, node_size=3000, node_color='blue', font_size=10, font_weight='bold', edge_color='gray', arrowsize=20)

# Save the plot to a temporary file
plt.savefig('graph.png')

# Add the plot to the Word document
doc.add_picture('graph.png', width=Inches(6))

# Save the Word document
doc.save('graph_document.docx')
