import matplotlib.pyplot as plt
import networkx as nx
import numpy as np
import csv
from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()
doc.add_heading('Integration Cybersecurity and Awareness Framework', level=1)

# Create your directed graph
direction = nx.DiGraph()

# Nodes for NIST construct
nist_nodes = ['Govern', 'Identity', 'Protect', 'Detect', 'Respond', 'Recover']

# Nodes for MediaPro contracts
mediapro_nodes = ['Analyze', 'Plan', 'Train', 'Reinforce']

# Adding the nodes to the graph
direction.add_nodes_from(nist_nodes)
direction.add_nodes_from(mediapro_nodes)

# Edges dependencies
nist_edges = [
    ('Govern', 'Identity'), ('Govern', 'Protect'), ('Govern', 'Detect'),
    ('Govern', 'Respond'), ('Govern', 'Recover'), ('Identity', 'Protect'),
    ('Protect', 'Detect'), ('Detect', 'Respond'), ('Respond', 'Recover'),
    ('Recover', 'Identity')
]

# MediaPro edges
mediaPro_edges = [
    ('Analyze', 'Plan'), ('Plan', 'Train'), ('Train', 'Reinforce'),
    ('Reinforce', 'Analyze')
]

# Integration edges
integrate_edges = [
    ('Analyze', 'Identity'), ('Plan', 'Protect'), ('Train', 'Detect'),
    ('Reinforce', 'Respond'), ('Analyze', 'Recover')
]

# Adding edges to the graph
direction.add_edges_from(nist_edges + mediaPro_edges + integrate_edges)

# Define positions
positions = {}

# Layers and their radius
layers = {
    1: {'radius': 0},
    2: {'radius': 3},
    3: {'radius': 6}
}

# Position for the center node "Govern"
center_x, center_y = 0, 0
positions['Govern'] = (center_x, center_y)

# Calculate positions for layer 2 nodes
num_layer2_nodes = len(nist_nodes) - 1  # Excluding 'Govern'
angle_step_layer2 = 2 * np.pi / num_layer2_nodes

for i, node in enumerate(nist_nodes):
    if node == 'Govern':
        continue  # Skip 'Govern' since its position is already set

    angle = i * angle_step_layer2
    x = center_x + layers[2]['radius'] * np.cos(angle)
    y = center_y + layers[2]['radius'] * np.sin(angle)
    positions[node] = (x, y)

# Place layer 3 nodes
num_layer3_nodes = len(mediapro_nodes)
angle_step_layer3 = 2 * np.pi / num_layer3_nodes

for i, node in enumerate(mediapro_nodes):
    angle = i * angle_step_layer3
    x = center_x + layers[3]['radius'] * np.cos(angle)
    y = center_y + layers[3]['radius'] * np.sin(angle)
    positions[node] = (x, y)

# Draw the graph
plt.figure(figsize=(10, 10))
nx.draw(direction,
        positions,
        with_labels=True,
        node_size=3000,
        node_color='blue',
        font_size=10,
        font_weight='bold',
        edge_color='gray',
        arrowsize=20,
        font_color='white',
        alpha=0.9)

# Save plot as image
plt.savefig('graph2.png')

# Add the image to the Word document
doc.add_picture('graph2.png', width=Inches(6))

# Save the Word document
doc.save('graph_do2c.docx')

# Show plot (optional)
plt.show()

# Export nodes to CSV with positions
with open('nodes_positions.csv', 'w', newline='') as node_file:
    writer = csv.writer(node_file)
    writer.writerow(['Id', 'Label', 'x', 'y'])
    for node, pos in positions.items():
        writer.writerow([node, node, pos[0], pos[1]])

# Export edges to CSV
with open('edges2.csv', 'w', newline='') as edge_file:
    writer = csv.writer(edge_file)
    writer.writerow(['Source', 'Target'])
    for edge in direction.edges():
        writer.writerow(edge)
